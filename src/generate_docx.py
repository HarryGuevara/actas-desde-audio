from docx import Document
from datetime import date

def crear_acta_docx(
    contenido: str,
    output_path: str = "data/outputs/ACTA_15.docx",
    fecha: str | None = None,
    asistentes: str = "…",
    titulo: str = "Acta de Reunión"
):
    if fecha is None:
        fecha = date.today().isoformat()
    doc = Document()
    doc.add_heading(titulo, level=1)
    doc.add_paragraph(f"Fecha: {fecha}")
    doc.add_paragraph(f"Asistentes: {asistentes}")
    doc.add_paragraph("\n--- Desarrollo de la reunión ---\n")
    doc.add_paragraph(contenido)
    doc.save(output_path)
    return output_path
